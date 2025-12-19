import streamlit as st
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io
import json

# Configure the OpenAI client with API key from Streamlit secrets
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# App title and description
st.title("Teacher's Assistant")
st.write("Upload multiple photos of a student's notebook pages (e.g., 3-5 pages covering a whole topic) to generate a comprehensive practice quiz.")

# Multi-file uploader for images
uploaded_files = st.file_uploader(
    "Upload notebook images",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True,
    help="Supported formats: JPG, PNG. Upload multiple pages covering a complete topic."
)

# Process the uploaded images
if uploaded_files:
    # Display upload count
    st.success(f"✓ {len(uploaded_files)} image(s) uploaded successfully")

    # Display all uploaded images in an expandable section
    with st.expander(f"View Uploaded Images ({len(uploaded_files)} pages)", expanded=False):
        cols = st.columns(min(len(uploaded_files), 3))  # Max 3 columns
        for idx, uploaded_file in enumerate(uploaded_files):
            col_idx = idx % 3
            with cols[col_idx]:
                st.image(uploaded_file, caption=f"Page {idx + 1}: {uploaded_file.name}", use_container_width=True)

    # Generate Practice Sheet button
    if st.button("Generate Practice Sheet", type="primary"):
        with st.spinner("Analyzing all notebook pages and generating comprehensive practice worksheet..."):

            # Prepare the content list with text prompt and all images
            content_list = [
                {
                    "type": "text",
                    "text": "Please analyze ALL of these notebook pages together as a complete set of notes on a topic. Detect the primary language of the handwritten notes and create a comprehensive practice quiz covering all the material across all pages. Output as pure JSON only."
                }
            ]

            # Add each image to the content list
            for uploaded_file in uploaded_files:
                # Reset file pointer to beginning
                uploaded_file.seek(0)

                # Convert image to Base64
                image_bytes = uploaded_file.getvalue()
                base64_image = base64.b64encode(image_bytes).decode("utf-8")

                # Determine the image MIME type
                file_extension = uploaded_file.name.split(".")[-1].lower()
                if file_extension in ["jpg", "jpeg"]:
                    mime_type = "image/jpeg"
                else:
                    mime_type = "image/png"

                # Add image to content list
                content_list.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{mime_type};base64,{base64_image}"
                    }
                })

            # System prompt for multi-page analysis with language detection
            system_prompt = """You are a helpful teacher's assistant specialized in analyzing student notebooks and creating practice materials.

You will be provided with MULTIPLE images of a student's handwritten notes that together cover a complete topic or lesson.

IMPORTANT INSTRUCTIONS:

1. MULTI-PAGE ANALYSIS: Analyze ALL provided images as a cohesive set of notes. The pages together represent complete coverage of a topic - synthesize the information across all pages.

2. LANGUAGE DETECTION: Detect the primary language of the handwritten notes (e.g., Spanish, English, French, Portuguese, etc.). Generate the ENTIRE practice worksheet in that SAME language. If the notes are in Spanish, ALL quiz content must be in Spanish. If in English, ALL content must be in English.

3. CONTENT ANALYSIS:
   - Identify the subject matter (e.g., Mathematics, Science, History, etc.)
   - Identify the specific topic spanning across the notes
   - Determine the appropriate grade level
   - Extract key concepts, definitions, formulas, dates, or facts from ALL pages

4. QUIZ GENERATION: Create a comprehensive 10-question multiple choice practice quiz that covers material from ALL the uploaded pages, not just one page.

5. OUTPUT FORMAT: You MUST output your response as a valid JSON object with NO markdown formatting around it (no ```json blocks). Use this EXACT structure:

{
  "subject": "The detected subject (e.g., Biology, Mathematics, History)",
  "topic": "The specific topic spanning the notes (e.g., Cell Division, Quadratic Equations, World War II)",
  "grade_level": "Detected or appropriate Grade Level",
  "questions": [
    {
      "q": "The question text here",
      "options": ["Option A", "Option B", "Option C", "Option D"],
      "answer": "The correct option text (must match one of the options exactly)"
    }
  ]
}

REQUIREMENTS:
- Generate exactly 10 questions that span content from ALL uploaded pages
- Each question must have exactly 4 options
- The "answer" field must contain the exact text of the correct option
- ALL text (subject, topic, grade_level, questions, options, answers) must be in the SAME language as the student's notes
- Use proper UTF-8 characters for special characters (ñ, á, é, í, ó, ú, ü, ¿, ¡, etc.)
- Questions should cover different concepts from across all the pages, not just repeat similar content
- Vary question difficulty to include both recall and comprehension questions"""

            # Make the API call with vision capabilities - all images in one request
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": content_list
                    }
                ],
                max_tokens=3000
            )

            # Extract the generated content
            raw_content = response.choices[0].message.content

            try:
                # Clean up response if it has markdown code blocks
                cleaned_content = raw_content.strip()
                if cleaned_content.startswith("```json"):
                    cleaned_content = cleaned_content[7:]
                if cleaned_content.startswith("```"):
                    cleaned_content = cleaned_content[3:]
                if cleaned_content.endswith("```"):
                    cleaned_content = cleaned_content[:-3]
                cleaned_content = cleaned_content.strip()

                # Parse the JSON response
                quiz_data = json.loads(cleaned_content)

                # Display the generated content in Streamlit
                st.subheader("Generated Practice Quiz")
                st.write(f"**Subject:** {quiz_data['subject']}")
                st.write(f"**Topic:** {quiz_data['topic']}")
                st.write(f"**Grade Level:** {quiz_data['grade_level']}")
                st.write(f"**Based on:** {len(uploaded_files)} notebook page(s)")
                st.write("---")

                for i, question in enumerate(quiz_data['questions'], 1):
                    st.write(f"**{i}. {question['q']}**")
                    for option in question['options']:
                        st.write(f"   • {option}")
                    st.write("")

                # Create professionally formatted Word document
                doc = Document()

                # Set document margins
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                # Add bold title using subject and topic
                title = doc.add_heading(level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.add_run(f"{quiz_data['subject']}: {quiz_data['topic']}")
                title_run.bold = True

                # Add grade level subtitle
                subtitle = doc.add_paragraph()
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_run = subtitle.add_run(f"{quiz_data['grade_level']} - Practice Quiz")
                subtitle_run.bold = True
                subtitle_run.font.size = Pt(14)

                # Add horizontal line effect with spacing
                doc.add_paragraph()
                separator = doc.add_paragraph("_" * 60)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

                # Iterate through questions
                for i, question in enumerate(quiz_data['questions'], 1):
                    # Add question text as bold paragraph
                    question_para = doc.add_paragraph()
                    question_run = question_para.add_run(f"{i}. {question['q']}")
                    question_run.bold = True
                    question_run.font.size = Pt(11)

                    # Add options as formal bullet points
                    option_labels = ['A', 'B', 'C', 'D']
                    for idx, option in enumerate(question['options']):
                        option_para = doc.add_paragraph(style='List Bullet')
                        option_text = f"{option_labels[idx]}) {option}"
                        option_run = option_para.add_run(option_text)
                        option_run.font.size = Pt(11)

                    # Add spacing between questions
                    doc.add_paragraph()

                # Add page break before answer key
                doc.add_page_break()

                # Add Answer Key section with bold heading
                answer_heading = doc.add_heading(level=1)
                answer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                answer_heading_run = answer_heading.add_run("Answer Key")
                answer_heading_run.bold = True

                doc.add_paragraph()

                # Add answers in a clean format
                for i, question in enumerate(quiz_data['questions'], 1):
                    # Find the correct option letter
                    correct_answer = question['answer']
                    option_labels = ['A', 'B', 'C', 'D']
                    correct_letter = ""
                    for idx, option in enumerate(question['options']):
                        if option == correct_answer:
                            correct_letter = option_labels[idx]
                            break

                    answer_para = doc.add_paragraph()
                    answer_para.add_run(f"{i}. ").bold = True
                    if correct_letter:
                        answer_para.add_run(f"{correct_letter}) {correct_answer}")
                    else:
                        answer_para.add_run(correct_answer)

                # Save document to in-memory buffer with UTF-8 encoding
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Create filename based on topic
                safe_topic = quiz_data['topic'].replace(" ", "_").replace("/", "-")[:30]
                filename = f"practice_quiz_{safe_topic}.docx"

                # Download button for the Word document
                st.download_button(
                    label="📥 Download Practice Quiz (Word Document)",
                    data=buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # Show success message
                st.success("Practice quiz generated successfully! Click the button above to download.")

            except json.JSONDecodeError as e:
                st.error("Failed to parse the AI response as JSON. Please try again.")
                st.text("Raw response:")
                st.code(raw_content)
            except KeyError as e:
                st.error(f"Missing expected field in response: {e}")
                st.text("Parsed data:")
                st.json(quiz_data)
else:
    # Show instructions when no files are uploaded
    st.info("👆 Upload one or more notebook page images to get started. For best results, upload all pages covering a complete topic.")
