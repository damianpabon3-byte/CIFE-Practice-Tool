"""
CIFE Edu-Suite - Exporter Module
=================================
Generates PDF and DOCX files for quiz worksheets and answer keys.
Creates professional, print-ready documents with proper formatting.

Supports enhanced JSON schema with:
- Schema version for forward compatibility
- Analysis results for re-generation
- Quiz settings for reproducibility
- Game state for resume functionality
"""

import io
import json
import re
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

# Schema version for JSON exports - increment when format changes
QUIZ_SCHEMA_VERSION = "2.0"


def _sanitize_text_for_pdf(text: str) -> str:
    """
    Sanitize text for PDF generation by replacing problematic unicode
    characters with ASCII equivalents or removing them.

    Args:
        text: Input text that may contain unicode

    Returns:
        Sanitized text safe for latin-1 encoding
    """
    if not text:
        return ""

    # Common unicode replacements
    replacements = {
        '\u2018': "'",   # Left single quote
        '\u2019': "'",   # Right single quote
        '\u201c': '"',   # Left double quote
        '\u201d': '"',   # Right double quote
        '\u2013': '-',   # En dash
        '\u2014': '--',  # Em dash
        '\u2026': '...', # Ellipsis
        '\u00a0': ' ',   # Non-breaking space
        '\u2022': '*',   # Bullet
        '\u00b7': '*',   # Middle dot
        '\u2212': '-',   # Minus sign
        '\u00d7': 'x',   # Multiplication sign
        '\u00f7': '/',   # Division sign
        '\u2264': '<=',  # Less than or equal
        '\u2265': '>=',  # Greater than or equal
        '\u2260': '!=',  # Not equal
        '\u00b0': ' deg',# Degree symbol
        '\u03c0': 'pi',  # Pi
        '\u00b2': '^2',  # Superscript 2
        '\u00b3': '^3',  # Superscript 3
        '\u221a': 'sqrt',# Square root
        '\u00bd': '1/2', # One half
        '\u00bc': '1/4', # One quarter
        '\u00be': '3/4', # Three quarters
    }

    for unicode_char, replacement in replacements.items():
        text = text.replace(unicode_char, replacement)

    # Remove any remaining non-latin-1 characters
    try:
        text.encode('latin-1')
    except UnicodeEncodeError:
        # Replace remaining problematic chars with ?
        text = text.encode('latin-1', errors='replace').decode('latin-1')

    return text


def _safe_multi_cell(pdf, w: float, h: float, text: str, max_chars: int = 500):
    """
    Safely render multi_cell with text length protection.
    Prevents "Not enough horizontal space" errors.

    Args:
        pdf: FPDF instance
        w: Cell width (0 = page width)
        h: Cell height
        text: Text to render
        max_chars: Maximum characters before truncation
    """
    # Truncate very long text to prevent layout issues
    if len(text) > max_chars:
        text = text[:max_chars] + "..."

    # Ensure we're at left margin
    pdf.set_x(10)

    # Use explicit width to prevent horizontal space errors
    effective_w = w if w > 0 else (pdf.w - pdf.l_margin - pdf.r_margin)

    try:
        pdf.multi_cell(effective_w, h, text)
    except Exception:
        # Fallback: try with smaller width
        try:
            pdf.multi_cell(effective_w - 10, h, text[:200] + "..." if len(text) > 200 else text)
        except Exception:
            # Ultimate fallback: use cell instead
            pdf.cell(effective_w - 10, h, text[:100] + "..." if len(text) > 100 else text, 0, 1)


def create_pdf(
    quiz_data: List[Dict[str, Any]],
    title: str = "Practice Quiz",
    subject: str = "",
    grade: str = "",
    include_answers: bool = False
) -> bytes:
    """
    Create a PDF document from quiz data.

    Generates a student worksheet (questions only) and optionally
    includes a teacher key with answers and explanations.

    Args:
        quiz_data: List of question dictionaries
        title: Title for the worksheet
        subject: Subject name to display
        grade: Grade level to display
        include_answers: Whether to include answer key

    Returns:
        PDF file as bytes

    Raises:
        Exception: If PDF generation fails
    """
    from fpdf import FPDF

    # Sanitize inputs early
    title = _sanitize_text_for_pdf(title or "Practice Quiz")
    subject = _sanitize_text_for_pdf(subject or "")
    grade_str = _sanitize_text_for_pdf(str(grade) if grade else "")

    class QuizPDF(FPDF):
        def __init__(self, doc_title: str, doc_subject: str, doc_grade: str):
            super().__init__()
            self.doc_title = doc_title
            self.doc_subject = doc_subject
            self.doc_grade = doc_grade
            self.set_auto_page_break(auto=True, margin=25)

        def header(self):
            self.set_font('Helvetica', 'B', 16)
            # Truncate title if too long
            display_title = self.doc_title[:60] + "..." if len(self.doc_title) > 60 else self.doc_title
            self.cell(0, 10, display_title, 0, 1, 'C')
            if self.doc_subject or self.doc_grade:
                self.set_font('Helvetica', '', 10)
                if self.doc_subject and self.doc_grade:
                    info = f"{self.doc_subject} - Grade {self.doc_grade}"
                else:
                    info = self.doc_subject or f"Grade {self.doc_grade}"
                self.cell(0, 6, info, 0, 1, 'C')
            self.set_font('Helvetica', 'I', 8)
            self.cell(0, 6, f"Generated: {datetime.now().strftime('%B %d, %Y')}", 0, 1, 'C')
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    pdf = QuizPDF(title, subject, grade_str)
    pdf.add_page()

    # Page width for calculations (A4 is 210mm, minus margins)
    left_margin = 10
    right_margin = 10
    effective_width = pdf.w - left_margin - right_margin  # ~190mm

    # Student Worksheet
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Student Worksheet', 0, 1, 'L')
    pdf.set_draw_color(79, 70, 229)
    pdf.line(left_margin, pdf.get_y(), pdf.w - right_margin, pdf.get_y())
    pdf.ln(5)

    # Name and Date fields
    pdf.set_font('Helvetica', '', 11)
    pdf.cell(95, 8, 'Name: _______________________', 0, 0, 'L')
    pdf.cell(95, 8, 'Date: ____________', 0, 1, 'R')
    pdf.ln(8)

    # Question type labels
    type_labels = {
        "multiple_choice": "[MC]",
        "true_false": "[T/F]",
        "short_answer": "[SA]",
        "matching": "[Match]",
        "fill_in_blank": "[Fill]"
    }

    # Questions
    for i, q in enumerate(quiz_data, 1):
        q_type = q.get("question_type", "multiple_choice")
        q_text = _sanitize_text_for_pdf(q.get("question_text", ""))
        badge = type_labels.get(q_type, "")

        # Question header
        pdf.set_font('Helvetica', 'B', 11)
        pdf.set_x(left_margin)
        _safe_multi_cell(pdf, effective_width, 7, f"{i}. {badge} {q_text}")

        # Options based on type
        pdf.set_font('Helvetica', '', 10)

        if q_type == "multiple_choice":
            options = q.get("options", [])
            labels = ["A", "B", "C", "D"]
            for j, opt in enumerate(options):
                if j < len(labels):
                    opt_text = _sanitize_text_for_pdf(str(opt))[:200]
                    pdf.set_x(left_margin + 10)
                    pdf.cell(effective_width - 10, 6, f"({labels[j]}) {opt_text}", 0, 1)
            pdf.ln(2)

        elif q_type == "true_false":
            pdf.set_x(left_margin + 10)
            pdf.cell(effective_width - 10, 6, "(   ) True    (   ) False", 0, 1)
            pdf.ln(2)

        elif q_type == "matching":
            # Matching: show left column items with blank for right column
            pairs = q.get("pairs", [])
            for j, pair in enumerate(pairs[:10]):  # Limit to 10 pairs
                left_item = _sanitize_text_for_pdf(str(pair.get("left", "")))[:50]
                pdf.set_x(left_margin + 10)
                pdf.cell(effective_width - 10, 6, f"{j+1}. {left_item} → _______", 0, 1)
            pdf.ln(2)

        elif q_type == "fill_in_blank":
            pdf.set_x(left_margin + 10)
            pdf.cell(effective_width - 10, 6, "Answer: _________________________________", 0, 1)
            pdf.ln(2)

        else:  # short_answer or other
            pdf.set_x(left_margin + 10)
            pdf.cell(effective_width - 10, 6, "Answer: _________________________________", 0, 1)
            pdf.ln(2)

        pdf.ln(3)

        # Check for page break with safe margin
        if pdf.get_y() > 245:
            pdf.add_page()

    # Teacher Answer Key (separate page)
    if include_answers:
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 14)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(0, 10, 'Teacher Answer Key', 0, 1, 'L')
        pdf.set_text_color(0, 0, 0)
        pdf.set_draw_color(79, 70, 229)
        pdf.line(left_margin, pdf.get_y(), pdf.w - right_margin, pdf.get_y())
        pdf.ln(8)

        for i, q in enumerate(quiz_data, 1):
            q_type = q.get("question_type", "multiple_choice")
            q_text = _sanitize_text_for_pdf(q.get("question_text", ""))
            correct = _sanitize_text_for_pdf(str(q.get("correct_answer", "")))
            explanation = _sanitize_text_for_pdf(q.get("explanation", ""))

            # Question
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_x(left_margin)
            _safe_multi_cell(pdf, effective_width, 6, f"{i}. {q_text}")

            # Correct answer
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(34, 197, 94)  # Green
            pdf.set_x(left_margin + 10)
            _safe_multi_cell(pdf, effective_width - 10, 6, f"Correct Answer: {correct}")
            pdf.set_text_color(0, 0, 0)

            # Explanation
            if explanation:
                pdf.set_font('Helvetica', 'I', 9)
                pdf.set_text_color(100, 100, 100)
                pdf.set_x(left_margin + 10)
                _safe_multi_cell(pdf, effective_width - 10, 5, f"Explanation: {explanation}")
                pdf.set_text_color(0, 0, 0)

            pdf.ln(4)

            if pdf.get_y() > 245:
                pdf.add_page()

    # Safe output conversion - handle bytearray, bytes, and str
    try:
        output = pdf.output(dest='S')
        if isinstance(output, bytearray):
            return bytes(output)
        if isinstance(output, str):
            return output.encode('latin-1', errors='ignore')
        if isinstance(output, bytes):
            return output
        # Fallback
        return bytes(output) if output else b''
    except Exception as e:
        raise Exception(f"PDF output conversion failed: {str(e)}")


def create_docx(
    quiz_data: List[Dict[str, Any]],
    title: str = "Practice Quiz",
    subject: str = "",
    grade: str = "",
    include_answers: bool = True
) -> bytes:
    """
    Create a Word document from quiz data.

    Creates an editable document with proper formatting.

    Args:
        quiz_data: List of question dictionaries
        title: Title for the worksheet
        subject: Subject name
        grade: Grade level
        include_answers: Whether to include answer key section

    Returns:
        DOCX file as bytes
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    if subject or grade:
        info = f"{subject} - Grade {grade}" if subject and grade else (subject or f"Grade {grade}")
        subtitle = doc.add_paragraph(info)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.italic = True
    date_para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Student section header
    section_title = doc.add_heading('Student Worksheet', level=1)
    section_title.runs[0].font.color.rgb = RGBColor(79, 70, 229)

    # Name/Date line
    name_date = doc.add_paragraph()
    name_date.add_run("Name: _________________________    Date: ____________")

    doc.add_paragraph()

    # Questions
    for i, q in enumerate(quiz_data, 1):
        q_type = q.get("question_type", "multiple_choice")
        q_text = q.get("question_text", "")

        type_labels = {
            "multiple_choice": "[MC]",
            "true_false": "[T/F]",
            "short_answer": "[SA]"
        }
        badge = type_labels.get(q_type, "")

        # Question paragraph
        q_para = doc.add_paragraph()
        q_num = q_para.add_run(f"{i}. {badge} ")
        q_num.bold = True
        q_para.add_run(q_text)

        # Options
        if q_type == "multiple_choice":
            options = q.get("options", [])
            labels = ["A", "B", "C", "D"]
            for j, opt in enumerate(options):
                if j < len(labels):
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Inches(0.5)
                    opt_para.add_run(f"({labels[j]}) {opt}")

        elif q_type == "true_false":
            tf_para = doc.add_paragraph()
            tf_para.paragraph_format.left_indent = Inches(0.5)
            tf_para.add_run("(   ) True    (   ) False")

        else:  # short_answer
            ans_para = doc.add_paragraph()
            ans_para.paragraph_format.left_indent = Inches(0.5)
            ans_para.add_run("Answer: _________________________________")

        doc.add_paragraph()

    # Answer Key Section
    if include_answers:
        doc.add_page_break()

        key_title = doc.add_heading('Teacher Answer Key', level=1)
        key_title.runs[0].font.color.rgb = RGBColor(79, 70, 229)

        doc.add_paragraph()

        for i, q in enumerate(quiz_data, 1):
            q_text = q.get("question_text", "")
            correct = q.get("correct_answer", "")
            explanation = q.get("explanation", "")

            # Question
            q_para = doc.add_paragraph()
            q_num = q_para.add_run(f"{i}. ")
            q_num.bold = True
            q_para.add_run(q_text)

            # Correct answer
            ans_para = doc.add_paragraph()
            ans_para.paragraph_format.left_indent = Inches(0.3)
            ans_label = ans_para.add_run("Correct Answer: ")
            ans_label.bold = True
            ans_value = ans_para.add_run(correct)
            ans_value.font.color.rgb = RGBColor(34, 197, 94)

            # Explanation
            if explanation:
                exp_para = doc.add_paragraph()
                exp_para.paragraph_format.left_indent = Inches(0.3)
                exp_label = exp_para.add_run("Explanation: ")
                exp_label.italic = True
                exp_text = exp_para.add_run(explanation)
                exp_text.font.color.rgb = RGBColor(100, 100, 100)
                exp_text.font.size = Pt(10)

            doc.add_paragraph()

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()


def create_json_export(
    quiz_data: List[Dict[str, Any]],
    metadata: Optional[Dict[str, Any]] = None,
    analysis_result: Optional[Dict[str, Any]] = None,
    quiz_settings: Optional[Dict[str, Any]] = None,
    game_state: Optional[Dict[str, Any]] = None
) -> str:
    """
    Create a JSON export of the quiz for sharing/importing.

    Enhanced schema v2.0 includes:
    - Schema version for forward compatibility
    - Analysis results for re-generation capability
    - Quiz settings for reproducibility
    - Game state for resume functionality

    Args:
        quiz_data: List of question dictionaries
        metadata: Optional metadata (title, subject, grade, language)
        analysis_result: Optional vision analysis results (transcribed_text, etc.)
        quiz_settings: Optional generation settings (counts, difficulty)
        game_state: Optional game state for resume (current_index, score, streak)

    Returns:
        JSON string with stable schema
    """
    # Build metadata with schema version
    full_metadata = metadata.copy() if metadata else {}
    full_metadata["schema_version"] = QUIZ_SCHEMA_VERSION

    export_data = {
        "schema_version": QUIZ_SCHEMA_VERSION,
        "created_at": datetime.now().isoformat(),
        "metadata": full_metadata,
        "questions": quiz_data
    }

    # Include optional sections if provided
    if analysis_result:
        # Only include safe analysis fields (no raw API responses)
        safe_analysis = {
            "transcribed_text": analysis_result.get("transcribed_text", ""),
            "subject": analysis_result.get("subject", ""),
            "detected_grade_level": analysis_result.get("detected_grade_level", ""),
            "core_concept": analysis_result.get("core_concept", ""),
            "language": analysis_result.get("language", ""),
            "key_terms": analysis_result.get("key_terms", []),
            "content_summary": analysis_result.get("content_summary", "")
        }
        export_data["analysis_result"] = safe_analysis

    if quiz_settings:
        export_data["quiz_settings"] = quiz_settings

    if game_state:
        export_data["game_state"] = game_state

    return json.dumps(export_data, indent=2, ensure_ascii=False)


def import_from_json(json_string: str) -> Tuple[List[Dict], Dict, Optional[Dict], Optional[Dict], Optional[Dict]]:
    """
    Import quiz data from a JSON string.

    Supports both v1.0 and v2.0 schema formats with graceful fallbacks.

    Args:
        json_string: JSON string containing quiz data

    Returns:
        Tuple of (questions, metadata, analysis_result, quiz_settings, game_state)
        - questions: List of question dictionaries
        - metadata: Dict with title, subject, grade, etc.
        - analysis_result: Optional vision analysis (v2.0+)
        - quiz_settings: Optional generation settings (v2.0+)
        - game_state: Optional game state for resume (v2.0+)

    Raises:
        ValueError: If JSON is invalid or format is unrecognized
    """
    try:
        data = json.loads(json_string)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON: {str(e)}")

    # Handle raw list format (oldest format)
    if isinstance(data, list):
        return data, {}, None, None, None

    # Handle dictionary formats (v1.0 and v2.0)
    if not isinstance(data, dict):
        raise ValueError("Invalid quiz format: expected object or array")

    # Extract questions (required)
    if "questions" in data:
        questions = data["questions"]
    else:
        raise ValueError("Invalid quiz format: missing 'questions' field")

    if not isinstance(questions, list):
        raise ValueError("Invalid quiz format: 'questions' must be an array")

    # Extract metadata with defaults
    metadata = data.get("metadata", {})
    if not isinstance(metadata, dict):
        metadata = {}

    # Add schema version to metadata if present at root
    if "schema_version" in data:
        metadata["schema_version"] = data["schema_version"]
    elif "version" in data:
        # v1.0 compatibility
        metadata["schema_version"] = data["version"]

    # Extract v2.0 fields (optional)
    analysis_result = data.get("analysis_result")
    quiz_settings = data.get("quiz_settings")
    game_state = data.get("game_state")

    # Validate and normalize questions
    normalized_questions = []
    for q in questions:
        if isinstance(q, dict) and q.get("question_text"):
            # Ensure required fields have defaults
            normalized_q = {
                "question_text": q.get("question_text", ""),
                "question_type": q.get("question_type", "multiple_choice"),
                "options": q.get("options", []),
                "correct_answer": q.get("correct_answer", ""),
                "correct_answer_index": q.get("correct_answer_index", -1),
                "explanation": q.get("explanation", ""),
                "misconception_tag": q.get("misconception_tag", "")
            }
            # Include any extra fields (like 'pairs' for matching)
            for key in q:
                if key not in normalized_q:
                    normalized_q[key] = q[key]
            normalized_questions.append(normalized_q)

    return normalized_questions, metadata, analysis_result, quiz_settings, game_state


def create_answer_sheet_pdf(
    quiz_data: List[Dict[str, Any]],
    title: str = "Answer Sheet"
) -> bytes:
    """
    Create a simple bubble-style answer sheet for scanning.

    Args:
        quiz_data: List of question dictionaries
        title: Title for the answer sheet

    Returns:
        PDF file as bytes
    """
    from fpdf import FPDF

    # Sanitize title
    title = _sanitize_text_for_pdf(title or "Answer Sheet")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)

    # Page width for calculations (A4 is 210mm, minus margins)
    left_margin = 10
    right_margin = 10
    effective_width = pdf.w - left_margin - right_margin

    # Title
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 10, title[:60], 0, 1, 'C')
    pdf.ln(5)

    # Name field
    pdf.set_font('Helvetica', '', 11)
    pdf.cell(0, 8, 'Name: _______________________________    Date: ____________', 0, 1, 'L')
    pdf.ln(10)

    # Instructions
    pdf.set_font('Helvetica', 'I', 9)
    pdf.set_x(left_margin)
    _safe_multi_cell(pdf, effective_width, 5,
        'Instructions: Fill in the circle next to your answer choice. For short answer questions, write your response on the line provided.')
    pdf.ln(8)

    pdf.set_font('Helvetica', '', 10)

    for i, q in enumerate(quiz_data, 1):
        q_type = q.get("question_type", "multiple_choice")

        pdf.set_x(left_margin)
        if q_type == "multiple_choice":
            pdf.cell(20, 8, f"{i}.", 0, 0, 'R')
            pdf.cell(effective_width - 20, 8, '(A)     (B)     (C)     (D)', 0, 1, 'L')

        elif q_type == "true_false":
            pdf.cell(20, 8, f"{i}.", 0, 0, 'R')
            pdf.cell(effective_width - 20, 8, '(T)     (F)', 0, 1, 'L')

        elif q_type == "matching":
            pdf.cell(20, 8, f"{i}.", 0, 0, 'R')
            pdf.cell(effective_width - 20, 8, '________', 0, 1, 'L')

        else:  # short_answer, fill_in_blank, or other
            pdf.cell(20, 8, f"{i}.", 0, 0, 'R')
            pdf.cell(effective_width - 20, 8, '________________________________', 0, 1, 'L')

        if pdf.get_y() > 265:
            pdf.add_page()

    # Safe output conversion
    try:
        output = pdf.output(dest='S')
        if isinstance(output, bytearray):
            return bytes(output)
        if isinstance(output, str):
            return output.encode('latin-1', errors='ignore')
        if isinstance(output, bytes):
            return output
        return bytes(output) if output else b''
    except Exception as e:
        raise Exception(f"PDF output conversion failed: {str(e)}")


def get_download_filename(title: str, extension: str) -> str:
    """
    Generate a safe filename for downloads.

    Args:
        title: Base title for the file
        extension: File extension (pdf, docx, json)

    Returns:
        Safe filename string
    """
    import re

    # Remove unsafe characters
    safe_title = re.sub(r'[^\w\s-]', '', title)
    safe_title = re.sub(r'[-\s]+', '_', safe_title)

    timestamp = datetime.now().strftime('%Y%m%d')

    return f"{safe_title}_{timestamp}.{extension}"
